from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── Helpers ───────────────────────────────────────────────────────────────

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)  # blue-800
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x1e, 0x3a, 0x8a)  # blue-900
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x37, 0x51, 0x8e)
    return p

def para(text="", bold_prefix=None, monospace=False):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        if monospace:
            r.font.name = "Courier New"
            r.font.size = Pt(9)
    if text:
        r = p.add_run(text)
        if monospace:
            r.font.name = "Courier New"
            r.font.size = Pt(9)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p

def code_block(lines):
    """Monospace indented block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    # light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F1F5F9")
    pPr.append(shd)
    return p

def table_2col(rows, col1_width=2.2, col2_width=4.1):
    """Simple 2-column table with header row."""
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    for i, (c1, c2) in enumerate(rows):
        t.cell(i, 0).width = Inches(col1_width)
        t.cell(i, 1).width = Inches(col2_width)
        r0 = t.cell(i, 0).paragraphs[0].add_run(c1)
        r1 = t.cell(i, 1).paragraphs[0].add_run(c2)
        r0.font.size = Pt(9)
        r1.font.size = Pt(9)
        if i == 0:
            r0.bold = True
            r1.bold = True
            # header shading
            for cell in (t.cell(i, 0), t.cell(i, 1)):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "DBEAFE")
                tcPr.append(shd)
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph("─" * 90)
    p.runs[0].font.color.rgb = RGBColor(0xcc, 0xd5, 0xe0)
    p.runs[0].font.size = Pt(7)

# ══════════════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════════════
t = doc.add_heading("Midterm Use Case Study Guide", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.runs[0].font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

sub = doc.add_paragraph("CSE416 — Redistricting Analysis Tool")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(11)
sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8b)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════
# UC 1 — GUI-17
# ══════════════════════════════════════════════════════════════════════════
heading("UC 1 — GUI-17: Display Box & Whisker Data  (SD)", level=1)
p = doc.add_paragraph()
p.add_run("Presenter: ").bold = True
p.add_run("Vivian Zhu")
doc.add_paragraph()

heading("Actor", 2)
bullet("End user (researcher / legislator) viewing the Ensemble Analysis section.")

heading("Preconditions", 2)
bullet("User has selected a state (AL or OR) and navigated to Ensemble Analysis.")
bullet("Backend has loaded the ensemble_analysis MongoDB collection via DataLoader.java.")
bullet("Box & whisker JSON was pre-computed offline (compute_boxwhisker.py) and stored in AL_data/.")

heading("Main Flow", 2)
steps = [
    'User clicks the "Box & Whisker" tab inside EnsembleAnalysisSection.jsx.',
    "React checks a useRef guard — first entry fires fetchEnsembleBoxWhisker(stateId) → GET /api/states/{stateId}/ensemble/box-whisker.",
    "Spring Boot EnsembleController reads EnsembleAnalysisDoc.boxWhisker from MongoDB (one document per state).",
    "Frontend receives bwData:",
]
for i, s in enumerate(steps, 1):
    bullet(f"{i}. {s}")

code_block([
    "bwData.feasibleGroups                    // ['black', 'white']",
    "bwData.ensembles[].ensembleType          // 'race-blind' | 'vra-constrained'",
    "bwData.ensembles[].groupDistricts[race]  // [{ index, min, q1, median, mean, q3, max }]",
    "bwData.enactedPlan.groupDistricts[race]  // [{ index, districtId, groupVapPercentage }]",
])

remaining = [
    "5. User picks a race via FeasibleRaceFilter sidebar → feasibleRaceFilter in Zustand updates.",
    "6. EnsembleAnalysisSection passes filtered arrays to two BoxWhiskerChart instances (race-blind + VRA side by side).",
    "7. BoxWhiskerChart.jsx renders a custom D3 SVG chart:",
]
for s in remaining:
    bullet(s)

bullet("X-axis: District index (ranked ascending by minority VAP %).", level=1)
bullet("Y-axis: Minority group VAP % (0–100%).", level=1)
bullet("Per district: IQR box (Q1→Q3), median line, dashed whiskers, amber Mean dot, pink Enacted dot.", level=1)
bullet("ResizeObserver hook keeps chart fully responsive.", level=1)
bullet("On hover: Tooltip div shows Max, Q3, Median, Mean, Q1, Min, Enacted.", level=1)
bullet("8. sharedYMax prop computed across both charts so Y-axes align for comparison.")

heading("Alternate Flows", 2)
bullet("No data: chart renders dashed placeholder with 'No data available.'")
bullet("Dots overlap (mean vs enacted within 10 px): they split ±5 px left/right.")
bullet("Tab revisit: useRef guard prevents re-fetch; existing bwData reused.")

heading("Postconditions", 2)
bullet("User can compare minority VAP% distributions (ensemble vs. enacted) for race-blind and VRA-constrained ensembles.")

heading("Key Technical Facts", 2)
table_2col([
    ("Concept", "Detail"),
    ("Chart library", "Custom D3 SVG — no Nivo (no Nivo box-whisker exists)"),
    ("D3 scales", "d3.scaleBand() for x (districts), d3.scaleLinear() for y (0→yMax)"),
    ("Box elements", "IQR rect, median line, dashed whiskers, cap lines, mean dot, enacted dot"),
    ("'District rank'", "Sorted by minority % within each plan; rank 1 = least-minority across all plans"),
    ("MongoDB doc", "One EnsembleAnalysisDoc per state; .boxWhisker holds both ensembles"),
    ("Endpoint", "GET /api/states/{stateId}/ensemble/box-whisker (lazy-fetched once)"),
    ("Data source", "export_boxwhisker_real_data.py reads compute_boxwhisker.py output → DataLoader"),
])

divider()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# UC 2 — Prepro-1
# ══════════════════════════════════════════════════════════════════════════
heading("UC 2 — Prepro-1: Integrate Multiple Data Sources  (AD)", level=1)
p = doc.add_paragraph()
p.add_run("Presenter: ").bold = True
p.add_run("Jason Yeung / Vivian Zhu")
doc.add_paragraph()

heading("Actor", 2)
bullet("Developer / system administrator running the preprocessing pipeline offline before app launch.")

heading("Preconditions", 2)
bullet("Raw files available: Census P.L. 94-171 block VAP CSV, TIGER block shapefiles, precinct GeoJSON (with enacted_cd), USDA RUCA CSV, ACS S1901 income CSV.")
bullet("Python environment with geopandas, pandas, gerrychain installed.")

heading("Data Sources Integrated", 2)
table_2col([
    ("Source", "What It Provides"),
    ("Census P.L. 94-171 CSV", "Block-level VAP by race (P4 table: total, HVAP, NH sub-groups)"),
    ("Census TIGER blocks (.shp)", "Block polygon geometries for spatial assignment"),
    ("Precinct GeoJSON", "Precinct polygons + election results (votes_dem/rep) + enacted_cd"),
    ("USDA RUCA CSV", "Urban/suburban/rural classification by Census tract code"),
    ("Census tracts shapefile", "Tract polygons for RUCA and income spatial joins"),
    ("ACS S1901 CSV", "Tract-level median & mean household income (S1901_C01_012E/013E)"),
], col1_width=2.5, col2_width=3.8)

heading("Main Flow  (mergingData.py)", 2)

p = doc.add_paragraph()
p.add_run("Step A — VAP: Census blocks → precincts").bold = True
steps_a = [
    "Read block VAP CSV; build 15-digit block GEOIDs (strip '1000000US', zero-pad).",
    "Locate NH sub-group columns via regex on P4 label row (e.g. 'Not Hispanic or Latino: White alone').",
    "Read TIGER block shapefile; merge VAP columns onto block geometries by GEOID_BLOCK.",
    "Spatial join (within): representative_point() of each block → sjoin(predicate='within') onto precincts. Deduplicate by GEOID_BLOCK. Unmatched blocks use sjoin_nearest() fallback.",
    "Aggregate (sum) VAP per precinct; merge back to precinct GeoDataFrame.",
    "State-specific race collapse:",
]
for i, s in enumerate(steps_a, 1):
    bullet(f"{i}. {s}")

code_block([
    "AL:  OTHER_VAP = VAP - NH_WHITE_ALONE_VAP - NH_BLACK_ALONE_VAP",
    "OR:  OTHER_VAP = VAP - NH_WHITE_ALONE_VAP - LATINO_VAP",
])

p = doc.add_paragraph()
p.add_run("Step B — RUCA region_type").bold = True
bullet("Read RUCA CSV; join onto tract shapefile by 11-digit TRACT_ID.")
bullet("Point-in-polygon: precinct rep point → tract polygon. Assign urban (1–3), suburban (4–6), rural (7–10).")

p = doc.add_paragraph()
p.add_run("Step C — ACS household income").bold = True
bullet("Read ACS S1901 CSV; join by TRACT_ID. Representative-point join: precinct → tract.")
bullet("Unmatched precincts: sjoin_nearest() fallback.")
bullet("Derives AVG_HH_INC (prefers mean, falls back to median). Imputes 0-income precincts from nearest non-zero-income tract (INCOME_IMPUTED=True).")

p = doc.add_paragraph()
p.add_run("Output:").bold = True
p.add_run(" AL_precincts_full.geojson — one feature per precinct with race VAP, region_type, income, election results, enacted_cd.")

p = doc.add_paragraph()
p.add_run("Post-processing pipeline:").bold = True
bullet("precinctGraph.py → builds GerryChain adjacency graph JSON (input to SeaWulf).")
bullet("export_heatmap_*_real_data.py → heatmap bin assignment JSONs.")
bullet("export_boxwhisker_real_data.py + compute_boxwhisker.py → box-whisker stats JSONs.")
bullet("DataLoader.java (--app.load-data=true) → seeds all MongoDB collections.")

heading("Alternate Flows", 2)
bullet("Unmatched blocks: sjoin_nearest() ensures no block is dropped.")
bullet("Missing income precincts: nearest-neighbor imputation (INCOME_IMPUTED=True).")
bullet("Duplicate block GEOIDs: groupby(GEOID_BLOCK).agg(sum) deduplication.")

heading("Postconditions", 2)
bullet("AL_precincts_full.geojson and OR_precincts_full.geojson contain all integrated attributes.")
bullet("Precinct adjacency graph JSON is ready for GerryChain / SeaWulf input.")

heading("Key Technical Facts", 2)
table_2col([
    ("Concept", "Detail"),
    ("Spatial join strategy", "representative_point() → sjoin(within) + sjoin_nearest() fallback"),
    ("Why representative_point()?", "Guaranteed inside the polygon; avoids vertex-on-boundary edge cases"),
    ("CRS for all joins", "EPSG:5070 — Albers Equal Area (meters), not lat/lon"),
    ("Why race VAP columns?", "GerryChain Tally updater needs one numeric column per group per node"),
    ("VAP not total pop", "Voting Age Population (18+) measures minority voting strength per Gingles / VRA"),
])

divider()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# UC 3 — SeaWulf-2
# ══════════════════════════════════════════════════════════════════════════
heading("UC 3 — SeaWulf-2: Run MGGG ReCom on SeaWulf  (AD)", level=1)
p = doc.add_paragraph()
p.add_run("Presenter: ").bold = True
p.add_run("Jason Yeung")
doc.add_paragraph()

heading("Actor", 2)
bullet("Developer submitting a SLURM batch job to SBU's SeaWulf HPC cluster.")

heading("Preconditions", 2)
bullet("AL_precincts_full.geojson (from Prepro-1) built and transferred to SeaWulf.")
bullet("precinctGraph.py converted it to GerryChain JSON (AL_graph.json).")
bullet("Conda env cse416-gis-py311 with gerrychain + geopandas available.")
bullet("Config file prepared: input/config_vra.json or input/config_raceblind.json.")

heading("SLURM Job Submission  (run_recom_final_vra.sbatch)", 2)
code_block([
    "#SBATCH -p debug-28core   # partition",
    "#SBATCH -N 1              # 1 node",
    "#SBATCH -c 8              # 8 CPU cores",
    "#SBATCH -t 00:30:00       # 30 min wall time",
    "#SBATCH -A class_cse416",
    "",
    "module load anaconda/3-new",
    "conda activate cse416-gis-py311",
    "python -u scripts/run_recom_parallel.py input/config_vra.json final",
])

heading("Config File Key Fields  (config_vra.json)", 2)
code_block([
    '{',
    '  "graph_path":       "input/AL_graph.json",',
    '  "assignment_col":   "enacted_cd",   // starting partition = enacted plan',
    '  "pop_col":          "VAP",',
    '  "num_districts":    7,',
    '  "pop_tolerance":    0.035,          // ±3.5% of ideal population',
    '  "steps_final":      5000,',
    '  "vra": {',
    '    "enabled":        true,',
    '    "group_key":      "NH_BLACK_ALONE_VAP",',
    '    "auto_thresholds": [0.50, 0.45, 0.40, 0.35, 0.30],',
    '    "effectiveness":  { "enabled": true, "party_of_choice": "D" }',
    '  }',
    '}',
])

heading("Main Flow", 2)
p = doc.add_paragraph()
p.add_run("Parallelization layer  (run_recom_parallel.py)").bold = True

parallel_steps = [
    "Reads SLURM_CPUS_PER_TASK=8 → launches 8 worker subprocesses.",
    "Splits steps_final=5000 evenly → 625 steps per worker (remainder to first workers).",
    "Each worker gets isolated config with its own output_dir under _parallel_tmp/worker_N/.",
    "subprocess.Popen() starts all 8 run_recom.py processes in parallel.",
    "Waits for all workers (p.wait()). Any failure → exits without merging.",
    "Concatenates JSONL outputs: plans_*.jsonl, boxwhisker_raw_*.jsonl, district_effectiveness_*.jsonl.",
    "Merges histograms (seat_splits, opp_hist, eff_hist, cut_hist) by summing counts across workers.",
    "Writes single merged summary_final.json.",
]
for i, s in enumerate(parallel_steps, 1):
    bullet(f"{i}. {s}")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Core ReCom chain  (run_recom.py)").bold = True

chain_steps = [
    "Loads Graph.from_json(graph_path) — each node is a precinct with VAP, race, and vote attributes.",
    "Sets up GerryChain Tally updaters: population, cut_edges, dem, rep, min_{group_key} per group.",
    "Creates initial = Partition(G, assignment='enacted_cd', updaters=...) — start = enacted plan.",
    "Computes ideal_pop = total_VAP / 7.",
]
for i, s in enumerate(chain_steps, 1):
    bullet(f"{i}. {s}")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("VRA constraint auto-detection:").bold = True
bullet("5. Walks thresholds [0.50, 0.45, 0.40, 0.35, 0.30]; picks highest where enacted plan has ≥1 opportunity district (minority% ≥ threshold). Sets chosen_thr.")
bullet("6. Adds vra_opp_constraint: any proposed plan must have ≥ target_k_opp opportunity districts.")
bullet("7. If effectiveness enabled: adds vra_eff_constraint: ≥ target_k_eff districts where minority% ≥ threshold AND party-of-choice (D) wins.")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("ReCom proposal:").bold = True
code_block([
    "proposal = partial(",
    "    recom,",
    "    pop_col='VAP', pop_target=ideal_pop, epsilon=0.035,",
    "    node_repeats=3,",
    "    method=partial(bipartition_tree, max_attempts=5000, allow_pair_reselection=True)",
    ")",
    "chain = MarkovChain(",
    "    proposal=proposal,",
    "    constraints=[pop_constraint, vra_opp_constraint, vra_eff_constraint],",
    "    accept=accept.always_accept,   # no rejection — pure random walk",
    "    initial_state=initial,",
    "    total_steps=5000,",
    ")",
])

p = doc.add_paragraph()
p.add_run("Per-step outputs  (written for every plan):").bold = True
bullet("plans_final.jsonl: {step, dem_seats, rep_seats, cut_edges, opp_districts, eff_districts, assignment?}")
bullet("boxwhisker_raw_final.jsonl: {step, group_key, threshold, district_pcts_sorted} — sorted minority% across 7 districts.")
bullet("district_effectiveness_final.jsonl: one row per (step, district, group_key) with is_effective, minority_pct, winner.")

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Post-SeaWulf processing:").bold = True
bullet("compute_boxwhisker.py: for each rank i (1–7), collects all plans' i-th sorted value → quantiles (min, Q1, median, Q3, max) → AL_boxwhisker_vra_black.json.")
bullet("Various export_*_real_data.py scripts build remaining MongoDB-ready JSONs.")
bullet("DataLoader.java seeds all MongoDB collections on server startup.")

heading("Alternate Flows", 2)
bullet("Worker failure: orchestrator detects non-zero exit code, skips merge, exits with error.")
bullet("Race-blind run: VRA constraints disabled (vra.enabled=false in config_raceblind.json); chain samples freely.")
bullet("VRA threshold never satisfied: falls back to lowest threshold (0.30).")

heading("Postconditions", 2)
bullet("plans_final.jsonl: 5000 redistricting plans.")
bullet("boxwhisker_raw_final.jsonl: per-plan sorted minority% vectors for all groups.")
bullet("summary_final.json: seat-split histogram, VRA opp/effectiveness histograms, cut-edge histogram.")
bullet("After post-processing: AL_boxwhisker_vra_black.json ready for MongoDB → frontend display.")

heading("Key Technical Facts", 2)
table_2col([
    ("Concept", "Detail"),
    ("Algorithm", "ReCom (Recombination): merge 2 adjacent districts, re-partition via spanning tree bipartition"),
    ("Library", "MGGG gerrychain — MarkovChain, Partition, recom, bipartition_tree"),
    ("Starting state", "Enacted plan (enacted_cd column from precinct GeoJSON)"),
    ("Population constraint", "within_percent_of_ideal_population(initial, 0.035) — ±3.5% of VAP/7"),
    ("VRA constraint", "Auto-detected threshold; plan must maintain ≥K opportunity districts"),
    ("Effectiveness constraint", "Opportunity district where party-of-choice (D) also wins the election"),
    ("Acceptance rule", "always_accept — pure random walk, no Metropolis-Hastings rejection"),
    ("Parallelism", "8 independent chains (different random seeds) merged by concatenation"),
    ("SeaWulf partition", "debug-28core, 1 node, 8 cores, 30-min wall time"),
    ("'District rank'", "Within each plan, districts sorted by minority% ascending; quantiles computed across plans at same rank position — not by district ID"),
    ("node_repeats=3", "Each bipartition may retry up to 3 times per node before resampling a new pair"),
])

divider()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# TEXT USE CASE — GINGLES
# ══════════════════════════════════════════════════════════════════════════
heading("Text Use Case — GUI: Display Gingles Analysis", level=1)
p = doc.add_paragraph()
p.add_run("Presenters: ").bold = True
p.add_run("Jason Yeung (SD), Vivian Zhu (AD) — applicable to all members")
doc.add_paragraph()

heading("Actor", 2)
bullet("End user (researcher / legislator) viewing the Racial Polarization section.")

heading("Preconditions", 2)
bullet("User has selected a state (AL or OR) and navigated to Racial Polarization.")
bullet("DataLoader.java has seeded the gingles MongoDB collection — one document per (stateId, race) pair (e.g. AL_black, AL_white).")
bullet("Gingles data was pre-computed offline by gingles.py from AL_precincts_full.geojson.")
bullet("Default feasibleRaceFilter in Zustand is 'black'.")

heading("Main Flow", 2)
bullet("1. User enters the Racial Polarization section. useActiveSection sets activeSection = 'racial-polarization' in Zustand. A one-time fetch for Vote/Seat Share fires immediately (GET /vote-seat-share) — gated at section level so the VS-SS tab can show enabled/disabled before the user clicks it.")
bullet("2. 'Gingles Analysis' tab is active by default. Three guards are checked: inRP === true, activeTab === 'gingles', and ginglesByRace['black'] not yet cached. All pass — fetch fires:")
code_block(["GET /api/states/AL/gingles?race=black"])
bullet("3. Spring Boot GinglesController queries MongoDB by _id = 'AL_black' and returns the full GinglesDoc:")
code_block([
    "points           — ~1,000–2,000 precincts: { id, name, x (minority VAP%), y (dem vote share%) }",
    "democraticTrendline  — OLS regression: [{ x, y }, ...]",
    "republicanTrendline  — OLS regression: [{ x, y }, ...]",
    "summaryRows      — binned summary: [{ rangeLabel, precinctCount, avgDemVoteShare, avgRepVoteShare }]",
])
bullet("4. Frontend caches result: ginglesByRace = { black: doc }. RacialPolarizationSection builds ginglesAdapted (useMemo) reshaping the cache into { feasibleSeriesByRace: { black: { points, trendlines, summaryRows } } }.")
bullet("5. GinglesScatterPlot renders using Nivo + D3 with four custom SVG layers:")
bullet("bgLayer — soft blue gradient background rectangle.", level=1)
bullet("thresholdLayer — dashed horizontal line at y=0.5 labeled '50% Majority' (key Gingles reference line).", level=1)
bullet("trendlineLayer — D3 monotone-X smoothed OLS trendlines for both parties, drawn with glow halo + solid stroke. A rising Democratic trendline as minority VAP% increases is evidence of racially polarized voting.", level=1)
bullet("customNodesLayer — replaces Nivo default nodes. Each precinct = two dots at same x (minority VAP%): blue at y=dem_share, red at y=1-dem_share. Selected dot enlarged (r=9).", level=1)
bullet("6. User clicks a dot. onDotClick(precinctId) fires → selectedId updates → dot enlarges on chart AND GinglesPrecinctTable scrolls to and highlights the matching row (precinct name, VAP%, vote shares, income, region type).")
bullet("7. User switches race to 'White' in FeasibleRaceFilter sidebar. feasibleRaceFilter = 'white'. ginglesByRace['white'] not cached → fires GET /api/states/AL/gingles?race=white. Once resolved, cached. Switching back to 'black' hits cache — no re-fetch.")

heading("Alternate Flows", 2)
bullet("No points: GinglesScatterPlot renders dashed placeholder ('No data available').")
bullet("Race already cached: ginglesByRace[race] guard fires and skips the fetch entirely.")
bullet("State change: all ginglesByRace cache entries cleared; all fetches restart fresh.")
bullet("VS-SS tab disabled: if voteSeatData.raciallyPolarized === false, tab is disabled with tooltip 'Gingles 2/3 not satisfied — racially polarized voting not detected in this state.'")

heading("Postconditions", 2)
bullet("User can inspect whether minority VAP% correlates with party vote share (Gingles precondition 2/3: racially polarized voting).")

heading("Key Technical Facts", 2)
table_2col([
    ("Concept", "Detail"),
    ("Why race is a query param", "Each race doc is ~500 KB. DataLoader splits the all-in-one source JSON into one MongoDB doc per race at load time. Query param selects which doc to fetch."),
    ("Cache strategy", "ginglesByRace map — once a race is fetched it is never re-fetched until state changes (cache-aside pattern)"),
    ("MongoDB doc key", "_id = '{stateId}_{race}' e.g. 'AL_black'"),
    ("Two dots per precinct", "Democratic dot at y=dem_share; Republican dot at y=1-dem_share. Same x (minority VAP%)."),
    ("Chart library", "@nivo/scatterplot for axes/nodes + D3 (d3.line, curveMonotoneX) for trendlines"),
    ("Trendline source", "Precomputed OLS regression in gingles_regression.py — not computed in browser"),
    ("Cross-component sync", "selectedId (local React state) passed as prop to both chart and table so clicking a dot highlights the table row"),
    ("Endpoint", "GET /api/states/{stateId}/gingles?race={race}"),
    ("Data source", "gingles.py reads AL_precincts_full.geojson → computes group_pct and dem_share per precinct → JSON loaded by DataLoader.java"),
])

divider()
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# TEXT USE CASE — PREPRO-4: STORE PREPROCESSED DATA
# ══════════════════════════════════════════════════════════════════════════
heading("Text Use Case — Prepro-4: Store Preprocessed Data  (AD)", level=1)
p = doc.add_paragraph()
p.add_run("Presenter: ").bold = True
p.add_run("Sumon Ahasan")
doc.add_paragraph()

heading("Actor", 2)
bullet("System administrator / developer running DataLoader.java to seed MongoDB from pre-computed JSON files.")

heading("Preconditions", 2)
bullet("All preprocessing scripts have run (Prepro-1 through Prepro-3, Prepro-7 through Prepro-10).")
bullet("The following JSON files exist in AL-real-data/ and OR-real-data/:")
bullet("AL-state-summary.json, AL-district-summary.json, AL-ensemble-summary.json", level=1)
bullet("AL-heatmap-precinct.json, AL-heatmap-census.json", level=1)
bullet("AL-Gingles-precinct.json", level=1)
bullet("AL-EI.json, AL-EI-compare.json", level=1)
bullet("AL-vote-seat-share.json", level=1)
bullet("Spring Boot backend is started with --app.load-data=true flag.")

heading("Main Flow", 2)
bullet("1. DataLoader.run() fires on startup. Drops all existing MongoDB collections (stateRepo, overviewRepo, heatmapRepo, ensembleRepo, ginglesRepo, eiKdeRepo, eiCompareRepo, vsRepo, geoRepo).")
bullet("2. loadStates() — reads splash-states.json → saves one StateDoc per state to states collection.")
bullet("3. loadGeoAssets() — reads US-48-States.geojson and AL/ORCongressionalDistricts.json → saves to geo_assets collection as-is.")
bullet("4. seedState('AL') runs the following in order:")

p = doc.add_paragraph()
p.add_run("4a. loadEiCompare() — first").bold = True
bullet("Reads AL-EI-compare.json → splits by race pair → saves one EiCompareDoc per pair (e.g. _id='AL_black_white') to ei_compare collection. Returns list of available pairs for embedding in overview.")

p = doc.add_paragraph()
p.add_run("4b. loadHeatmaps() — second").bold = True
bullet("Reads AL-heatmap-precinct.json and AL-heatmap-census.json.")
bullet("Source format has all races per feature: { idx, black: binId, white: binId, ... }")
bullet("DataLoader discovers races dynamically from the first feature's key set (every key except 'idx').")
bullet("For each race x granularity: slims features to { idx, binId } for that race only → saves one HeatmapDoc.")
bullet("Example output docs: AL_precinct_black, AL_precinct_white, AL_census_block_black, AL_census_block_white.")
bullet("Returns list of discovered race keys for embedding in overview.")

p = doc.add_paragraph()
p.add_run("4c. loadOverview()").bold = True
bullet("Reads AL-state-summary.json + AL-district-summary.json + AL-ensemble-summary.json.")
bullet("Combines all three into one StateOverviewDoc (_id='AL') + embeds available heatmap races and EI compare pairs.")
bullet("Saves to state_overview collection.")

p = doc.add_paragraph()
p.add_run("4d. loadEnsemble()").bold = True
bullet("Reads AL-splits.json and AL-boxwhisker.json.")
bullet("Combines into one EnsembleAnalysisDoc (_id='AL') with .splits and .boxWhisker fields.")
bullet("Saves to ensemble_analysis collection.")

p = doc.add_paragraph()
p.add_run("4e. loadGingles()").bold = True
bullet("Reads AL-Gingles-precinct.json — source has all races in one file under feasibleSeriesByRace map.")
bullet("Iterates the map → for each race creates a separate GinglesDoc (_id='AL_black', 'AL_white', etc.).")
bullet("Each doc gets: points, democraticTrendline, republicanTrendline, summaryRows.")
bullet("Saves to gingles collection.")

p = doc.add_paragraph()
p.add_run("4f. loadEiKde()").bold = True
bullet("Reads AL-EI.json → splits by race → saves one EiKdeDoc per race to ei_kde collection.")

p = doc.add_paragraph()
p.add_run("4g. loadVoteSeatShare()").bold = True
bullet("Reads AL-vote-seat-share.json → saves one VoteSeatShareDoc (_id='AL') to vote_seat_share collection as-is.")

bullet("5. Steps 4a–4g repeat for OR.")
bullet("6. DataLoader prints '[DataLoader] Seed complete.' App restarts normally with load-data=false.")

heading("Alternate Flows", 2)
bullet("File not found: each loader checks readIfExists() — if null, prints a skip message and continues. Missing file does not crash the loader.")
bullet("Re-seeding: all collections are dropped at the start, so re-running with load-data=true always produces a clean state.")

heading("Postconditions", 2)
bullet("All MongoDB collections are fully populated and ready to serve API requests.")
bullet("No further preprocessing is needed until new data is added.")

heading("Key Technical Facts", 2)
table_2col([
    ("Concept", "Detail"),
    ("Trigger", "Spring Boot CommandLineRunner — fires once on startup when app.load-data=true"),
    ("Drop-then-reseed", "All collections deleted before seeding — guarantees clean consistent state"),
    ("Load time vs query time", "DataLoader splits source files into per-race/per-granularity docs at load time so each API query is a single MongoDB document lookup by _id — no filtering at query time"),
    ("Heatmap split", "Source: all races per feature. Output: one doc per race x granularity with only { idx, binId }"),
    ("Gingles split", "Source: all races under feasibleSeriesByRace. Output: one doc per race"),
    ("Ensemble merged", "Two source files (splits + boxwhisker) combined into one doc — served by two separate endpoints"),
    ("Overview embedded", "Available heatmap races and EI compare pairs discovered at load time and embedded in StateOverviewDoc so frontend knows valid query options"),
    ("Precinct GeoJSONs NOT loaded", "Too large for MongoDB 16 MB limit — streamed from disk by PrecinctGeoController instead"),
])

divider()

# ══════════════════════════════════════════════════════════════════════════
# END-TO-END FLOW
# ══════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("End-to-End System Flow", level=1)

code_block([
    "Raw Census CSVs + Shapefiles + Election Results",
    "        ↓  mergingData.py  (Prepro-1)",
    "AL_precincts_full.geojson",
    "        ↓  precinctGraph.py",
    "AL_graph.json  (GerryChain precinct adjacency)",
    "        ↓  sbatch run_recom_final_vra.sbatch  (SeaWulf-2)",
    "plans_final.jsonl  +  boxwhisker_raw_final.jsonl",
    "        ↓  compute_boxwhisker.py",
    "AL_boxwhisker_vra_black.json  (quantile stats per district rank)",
    "        ↓  DataLoader.java  --app.load-data=true",
    "MongoDB: ensemble_analysis collection  (EnsembleAnalysisDoc)",
    "        ↓  GET /api/states/AL/ensemble/box-whisker",
    "EnsembleAnalysisSection.jsx  (React frontend)",
    "        ↓  BoxWhiskerChart.jsx  (GUI-17)",
    "D3 SVG box-and-whisker chart in browser",
])

doc.add_paragraph()
heading("Quick-Reference: Who Does What", level=2)
table_2col([
    ("Use Case", "Key Responsibility"),
    ("Prepro-1 (Jason/Vivian)", "Build enriched precinct GeoJSON integrating VAP, RUCA, income onto precinct polygons"),
    ("SeaWulf-2 (Jason)", "Run parallel ReCom Markov chains on HPC; produce 5000 sampled redistricting plans"),
    ("GUI-17 (Vivian)", "Render responsive D3 SVG box-and-whisker chart comparing ensemble vs. enacted plan"),
])

# ── Save ──────────────────────────────────────────────────────────────────
out = "c:/Users/T14s/CSE416/Midterm_Use_Case_Study_Guide.docx"
doc.save(out)
print(f"Saved: {out}")
